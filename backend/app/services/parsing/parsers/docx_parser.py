from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.schemas.doc_parsing import ParseResult
from app.services.parsing.base import BaseDocumentParser
from app.services.parsing.cleaner import PageContent, TextCleaner
from app.services.parsing.exception import (
    EmptyDocumentError,
    ParserExecutionError,
)


class DocxDocumentParser(BaseDocumentParser):
    supported_extensions = frozenset({".docx"})

    supported_mime_types = frozenset(
        {
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document",
        }
    )

    def __init__(self, cleaner: TextCleaner) -> None:
        self._cleaner = cleaner

    def parse(
        self,
        file_path: Path,
        mime_type: str | None = None,
    ) -> ParseResult:
        try:
            document = DocxDocument(str(file_path))
        except Exception as exc:
            raise ParserExecutionError(
                f"Unable to open DOCX document: {exc}"
            ) from exc

        blocks: list[str] = []
        table_count = 0

        try:
            for item in document.iter_inner_content():
                if isinstance(item, Paragraph):
                    paragraph_text = item.text.strip()

                    if paragraph_text:
                        blocks.append(paragraph_text)

                elif isinstance(item, Table):
                    table_count += 1

                    table_markdown = self._table_to_markdown(item)

                    if table_markdown:
                        blocks.append(table_markdown)

        except AttributeError:
            # Compatibility fallback for older python-docx versions.
            for paragraph in document.paragraphs:
                paragraph_text = paragraph.text.strip()

                if paragraph_text:
                    blocks.append(paragraph_text)

            for table in document.tables:
                table_count += 1

                table_markdown = self._table_to_markdown(table)

                if table_markdown:
                    blocks.append(table_markdown)

        raw_text = "\n\n".join(blocks)

        cleaned_text, warnings = self._cleaner.clean_pages(
            [
                PageContent(
                    page_number=1,
                    text=raw_text,
                )
            ]
        )

        if not cleaned_text.strip():
            raise EmptyDocumentError(
                "The DOCX document contains no extractable text."
            )

        return ParseResult(
            text=cleaned_text,
            page_count=1,
            character_count=len(cleaned_text),
            detected_type="docx",
            detected_encoding=None,
            used_ocr=False,
            ocr_page_count=0,
            table_count=table_count,
            warnings=warnings,
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": table_count,
            },
        )

    @staticmethod
    def _table_to_markdown(table: Table) -> str:
        rows: list[list[str]] = []

        for row in table.rows:
            values = [
                cell.text.replace("\n", " ").strip()
                for cell in row.cells
            ]

            if any(values):
                rows.append(values)

        if not rows:
            return ""

        maximum_columns = max(len(row) for row in rows)

        normalized_rows = [
            row + [""] * (maximum_columns - len(row))
            for row in rows
        ]

        header = normalized_rows[0]
        separator = ["---"] * maximum_columns

        markdown_lines = [
            DocxDocumentParser._markdown_row(header),
            DocxDocumentParser._markdown_row(separator),
        ]

        markdown_lines.extend(
            DocxDocumentParser._markdown_row(row)
            for row in normalized_rows[1:]
        )

        return "\n".join(markdown_lines)

    @staticmethod
    def _markdown_row(values: list[str]) -> str:
        escaped_values = [
            value.replace("|", "\\|")
            for value in values
        ]

        return f"| {' | '.join(escaped_values)} |"
    